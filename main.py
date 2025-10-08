from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file
from werkzeug.security import generate_password_hash, check_password_hash
import sqlite3, os, json, requests, io
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
from datetime import datetime, timezone, timedelta
from openai import OpenAI
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import stripe

app = Flask(__name__)
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", "dev-secret-change-me")
app.config["PREFERRED_URL_SCHEME"] = "https"
# Session cookie configuration for OAuth redirects
app.config["SESSION_COOKIE_SAMESITE"] = "None"
app.config["SESSION_COOKIE_SECURE"] = True
app.config["SESSION_COOKIE_HTTPONLY"] = True
DB_PATH = os.environ.get("DB_PATH", "smartoffice.db")

# Stripe configuration
STRIPE_SECRET_KEY = os.environ.get("STRIPE_SECRET_KEY")
if STRIPE_SECRET_KEY:
    stripe.api_key = STRIPE_SECRET_KEY
else:
    print("WARNING: STRIPE_SECRET_KEY not set - subscription features will be disabled")

# OAuth configuration - Must be set as environment variables
GOOGLE_CLIENT_ID = os.environ.get("GOOGLE_CLIENT_ID")
GOOGLE_CLIENT_SECRET = os.environ.get("GOOGLE_CLIENT_SECRET")
SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']

if not GOOGLE_CLIENT_ID or not GOOGLE_CLIENT_SECRET:
    print("WARNING: GOOGLE_CLIENT_ID and GOOGLE_CLIENT_SECRET environment variables must be set for Gmail integration")

# OpenAI configuration
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
openai_client = None
if OPENAI_API_KEY:
    openai_client = OpenAI(api_key=OPENAI_API_KEY)
else:
    print("INFO: OPENAI_API_KEY not set - AI features will be disabled")

# Translations
TRANSLATIONS = {
    'en': {
        'app_name': 'SmartOffice AI',
        'email': 'Email',
        'calendar': 'Calendar',
        'reports': 'Reports',
        'logout': 'Logout',
        'login': 'Login',
        'welcome': 'Welcome',
        'dashboard': 'Dashboard',
        'name': 'Name',
        'password': 'Password',
        'register': 'Register',
        'dont_have_account': "Don't have an account?",
        'already_have_account': 'Already have an account?',
        'sign_up': 'Sign Up',
        'sign_in': 'Sign in',
        'gmail_inbox': 'Gmail Inbox',
        'messages': 'messages',
        'connected': 'Connected',
        'not_connected': 'Not Connected',
        'urgent': 'Urgent',
        'important': 'Important',
        'normal': 'Normal',
        'upcoming_events': 'Upcoming Events',
        'events': 'events',
        'generate_report': 'Generate Report',
        'report_period': 'Report Period',
        'past_week': 'Past Week',
        'past_month': 'Past Month',
        'custom_range': 'Custom Date Range',
        'pdf_doc': 'PDF Document',
        'word_doc': 'Word Document',
        'include_in_report': 'Include in Report',
        'email_summary': 'Email Summary',
        'calendar_events': 'Calendar Events',
        'upgrade_to_premium': 'Upgrade to Premium',
        'premium_active': 'Premium Active',
        'free_plan': 'You\'re currently on the Free plan',
        'upgrade_message': 'Upgrade to access Email, Calendar, Reports, and AI-powered features.',
        'view_plans': 'View Plans',
        'full_access': 'You have full access to all SmartOffice AI features.',
        'manage_inbox': 'Manage your Gmail inbox with AI-powered organization',
        'stay_on_schedule': 'Stay on top of your schedule and meetings',
        'generate_reports': 'Generate automated weekly reports',
        'view_inbox': 'View Inbox',
        'view_calendar': 'View Calendar',
        'create_report': 'Create Report',
        'premium_only': 'Premium Only',
        'choose_plan': 'Choose Your Plan',
        'unlock_power': 'Unlock the full power of SmartOffice AI with our subscription',
        'free': 'Free',
        'premium': 'Premium',
        'per_month': '/month',
        'current_plan': 'Current Plan',
        'subscribe_now': 'Subscribe Now',
        'recommended': 'RECOMMENDED',
        'basic_dashboard': 'Basic dashboard access',
        'email_management': 'Email management',
        'calendar_integration': 'Calendar integration',
        'report_generation': 'Report generation',
        'ai_features': 'AI-powered features',
        'full_dashboard': 'Full dashboard access',
        'gmail_inbox_mgmt': 'Gmail inbox management',
        'google_calendar': 'Google Calendar integration',
        'pdf_word_reports': 'PDF/Word report generation',
        'ai_email_priority': 'AI email prioritization',
        'ai_reply_suggest': 'AI reply suggestions',
        'bilingual_support': 'Bilingual support (EN/AR)',
        'email_support': '24/7 email support',
        'secure_payment': 'Secure payment powered by Stripe',
        'cancel_anytime': 'Cancel anytime',
        'payment_successful': 'Payment Successful!',
        'welcome_premium': 'Welcome to SmartOffice AI Premium! Your subscription is now active.',
        'you_have_access': 'You now have access to:',
        'go_to_dashboard': 'Go to Dashboard',
        'payment_cancelled': 'Payment Cancelled',
        'payment_cancel_msg': 'Your subscription payment was cancelled. No charges were made.',
        'try_again': 'You can try again anytime to unlock premium features.',
        'back_to_pricing': 'Back to Pricing',
        'reports_description': 'Generate weekly/monthly activity reports from your emails and calendar'
    },
    'ar': {
        'app_name': 'سمارت أوفيس AI',
        'email': 'البريد الإلكتروني',
        'calendar': 'التقويم',
        'reports': 'التقارير',
        'logout': 'تسجيل الخروج',
        'login': 'تسجيل الدخول',
        'welcome': 'مرحباً',
        'dashboard': 'لوحة التحكم',
        'name': 'الاسم',
        'password': 'كلمة المرور',
        'register': 'تسجيل',
        'dont_have_account': 'ليس لديك حساب؟',
        'already_have_account': 'لديك حساب بالفعل؟',
        'sign_up': 'إنشاء حساب',
        'sign_in': 'تسجيل الدخول',
        'gmail_inbox': 'صندوق وارد Gmail',
        'messages': 'رسائل',
        'connected': 'متصل',
        'not_connected': 'غير متصل',
        'urgent': 'عاجل',
        'important': 'مهم',
        'normal': 'عادي',
        'upcoming_events': 'الأحداث القادمة',
        'events': 'أحداث',
        'generate_report': 'إنشاء تقرير',
        'report_period': 'فترة التقرير',
        'past_week': 'الأسبوع الماضي',
        'past_month': 'الشهر الماضي',
        'custom_range': 'نطاق مخصص',
        'pdf_doc': 'مستند PDF',
        'word_doc': 'مستند Word',
        'include_in_report': 'تضمين في التقرير',
        'email_summary': 'ملخص البريد',
        'calendar_events': 'أحداث التقويم',
        'upgrade_to_premium': 'الترقية إلى بريميوم',
        'premium_active': 'الاشتراك المميز نشط',
        'free_plan': 'أنت حالياً على الخطة المجانية',
        'upgrade_message': 'قم بالترقية للوصول إلى البريد الإلكتروني والتقويم والتقارير والميزات المدعومة بالذكاء الاصطناعي.',
        'view_plans': 'عرض الخطط',
        'full_access': 'لديك وصول كامل إلى جميع ميزات سمارت أوفيس AI.',
        'manage_inbox': 'إدارة صندوق وارد Gmail الخاص بك مع تنظيم مدعوم بالذكاء الاصطناعي',
        'stay_on_schedule': 'ابق على اطلاع على جدولك واجتماعاتك',
        'generate_reports': 'إنشاء تقارير أسبوعية تلقائية',
        'view_inbox': 'عرض صندوق الوارد',
        'view_calendar': 'عرض التقويم',
        'create_report': 'إنشاء تقرير',
        'premium_only': 'للمشتركين المميزين فقط',
        'choose_plan': 'اختر خطتك',
        'unlock_power': 'اكتشف القوة الكاملة لسمارت أوفيس AI مع اشتراكنا',
        'free': 'مجاني',
        'premium': 'بريميوم',
        'per_month': '/شهرياً',
        'current_plan': 'الخطة الحالية',
        'subscribe_now': 'اشترك الآن',
        'recommended': 'موصى به',
        'basic_dashboard': 'وصول أساسي للوحة التحكم',
        'email_management': 'إدارة البريد الإلكتروني',
        'calendar_integration': 'تكامل التقويم',
        'report_generation': 'إنشاء التقارير',
        'ai_features': 'ميزات الذكاء الاصطناعي',
        'full_dashboard': 'وصول كامل للوحة التحكم',
        'gmail_inbox_mgmt': 'إدارة صندوق وارد Gmail',
        'google_calendar': 'تكامل تقويم Google',
        'pdf_word_reports': 'إنشاء تقارير PDF/Word',
        'ai_email_priority': 'ترتيب أولوية البريد بالذكاء الاصطناعي',
        'ai_reply_suggest': 'اقتراحات الرد بالذكاء الاصطناعي',
        'bilingual_support': 'دعم ثنائي اللغة (إنجليزي/عربي)',
        'email_support': 'دعم على مدار الساعة',
        'secure_payment': 'دفع آمن مدعوم من Stripe',
        'cancel_anytime': 'إلغاء في أي وقت',
        'payment_successful': 'تم الدفع بنجاح!',
        'welcome_premium': 'مرحباً بك في سمارت أوفيس AI بريميوم! اشتراكك نشط الآن.',
        'you_have_access': 'لديك الآن وصول إلى:',
        'go_to_dashboard': 'الذهاب إلى لوحة التحكم',
        'payment_cancelled': 'تم إلغاء الدفع',
        'payment_cancel_msg': 'تم إلغاء دفع اشتراكك. لم يتم فرض أي رسوم.',
        'try_again': 'يمكنك المحاولة مرة أخرى في أي وقت لفتح الميزات المميزة.',
        'back_to_pricing': 'العودة إلى الأسعار',
        'reports_description': 'إنشاء تقارير نشاط أسبوعية/شهرية من رسائلك البريدية والتقويم'
    }
}

def get_language():
    """Get current language from session, default to English"""
    return session.get('language', 'en')

def t(key):
    """Translate a key to the current language"""
    lang = get_language()
    return TRANSLATIONS.get(lang, TRANSLATIONS['en']).get(key, key)

@app.context_processor
def inject_language():
    """Make translation function and language available in all templates"""
    return dict(t=t, lang=get_language())

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db()
    conn.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT UNIQUE NOT NULL,
            name TEXT NOT NULL,
            password_hash TEXT NOT NULL,
            subscription_status TEXT DEFAULT 'free',
            stripe_customer_id TEXT,
            subscription_end_date TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS gmail_tokens (
            user_id INTEGER PRIMARY KEY,
            token TEXT NOT NULL,
            refresh_token TEXT,
            token_uri TEXT,
            token_expiry TEXT,
            scopes TEXT,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS calendar_tokens (
            user_id INTEGER PRIMARY KEY,
            token TEXT NOT NULL,
            refresh_token TEXT,
            token_uri TEXT,
            token_expiry TEXT,
            scopes TEXT,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    """)
    
    # Add subscription columns to existing users if they don't exist
    try:
        conn.execute("ALTER TABLE users ADD COLUMN subscription_status TEXT DEFAULT 'free'")
    except sqlite3.OperationalError:
        pass
    try:
        conn.execute("ALTER TABLE users ADD COLUMN stripe_customer_id TEXT")
    except sqlite3.OperationalError:
        pass
    try:
        conn.execute("ALTER TABLE users ADD COLUMN subscription_end_date TEXT")
    except sqlite3.OperationalError:
        pass
    
    conn.commit()
    conn.close()

with app.app_context():
    init_db()

def get_user_gmail_credentials(user_id):
    """Get Gmail credentials from database for a specific user"""
    conn = get_db()
    token_row = conn.execute("SELECT * FROM gmail_tokens WHERE user_id = ?", (user_id,)).fetchone()
    conn.close()
    
    if not token_row:
        return None
    
    from datetime import datetime
    expiry = None
    if token_row['token_expiry']:
        try:
            expiry = datetime.fromisoformat(token_row['token_expiry'])
        except:
            pass
    
    creds = Credentials(
        token=token_row['token'],
        refresh_token=token_row['refresh_token'],
        token_uri=token_row['token_uri'],
        client_id=GOOGLE_CLIENT_ID,
        client_secret=GOOGLE_CLIENT_SECRET,
        scopes=json.loads(token_row['scopes']) if token_row['scopes'] else [],
        expiry=expiry
    )
    
    # Refresh token if expired
    if creds.expired and creds.refresh_token:
        from google.auth.transport.requests import Request
        creds.refresh(Request())
        save_user_gmail_credentials(user_id, creds)
    
    return creds

def save_user_gmail_credentials(user_id, credentials):
    """Save Gmail credentials to database for a specific user"""
    conn = get_db()
    
    expiry_str = None
    if credentials.expiry:
        expiry_str = credentials.expiry.isoformat()
    
    refresh_token = credentials.refresh_token
    if not refresh_token:
        existing = conn.execute(
            "SELECT refresh_token FROM gmail_tokens WHERE user_id = ?", 
            (user_id,)
        ).fetchone()
        if existing:
            refresh_token = existing['refresh_token']
    
    conn.execute("""
        INSERT OR REPLACE INTO gmail_tokens 
        (user_id, token, refresh_token, token_uri, token_expiry, scopes)
        VALUES (?, ?, ?, ?, ?, ?)
    """, (
        user_id,
        credentials.token,
        refresh_token,
        credentials.token_uri,
        expiry_str,
        json.dumps(credentials.scopes) if credentials.scopes else '[]'
    ))
    conn.commit()
    conn.close()

def get_gmail_service(user_id):
    """Get authenticated Gmail service for a specific user"""
    creds = get_user_gmail_credentials(user_id)
    if not creds:
        return None
    
    try:
        service = build('gmail', 'v1', credentials=creds)
        return service
    except Exception as e:
        print(f"Error building Gmail service: {e}")
        return None

def get_user_calendar_credentials(user_id):
    """Get Calendar credentials from database for a specific user"""
    conn = get_db()
    token_row = conn.execute("SELECT * FROM calendar_tokens WHERE user_id = ?", (user_id,)).fetchone()
    conn.close()
    
    if not token_row:
        return None
    
    from datetime import datetime
    expiry = None
    if token_row['token_expiry']:
        try:
            expiry = datetime.fromisoformat(token_row['token_expiry'])
        except:
            pass
    
    creds = Credentials(
        token=token_row['token'],
        refresh_token=token_row['refresh_token'],
        token_uri=token_row['token_uri'],
        client_id=GOOGLE_CLIENT_ID,
        client_secret=GOOGLE_CLIENT_SECRET,
        scopes=json.loads(token_row['scopes']) if token_row['scopes'] else [],
        expiry=expiry
    )
    
    # Refresh token if expired
    if creds.expired and creds.refresh_token:
        from google.auth.transport.requests import Request
        creds.refresh(Request())
        save_user_calendar_credentials(user_id, creds)
    
    return creds

def save_user_calendar_credentials(user_id, credentials):
    """Save Calendar credentials to database for a specific user"""
    conn = get_db()
    
    expiry_str = None
    if credentials.expiry:
        expiry_str = credentials.expiry.isoformat()
    
    refresh_token = credentials.refresh_token
    if not refresh_token:
        existing = conn.execute(
            "SELECT refresh_token FROM calendar_tokens WHERE user_id = ?", 
            (user_id,)
        ).fetchone()
        if existing:
            refresh_token = existing['refresh_token']
    
    conn.execute("""
        INSERT OR REPLACE INTO calendar_tokens 
        (user_id, token, refresh_token, token_uri, token_expiry, scopes)
        VALUES (?, ?, ?, ?, ?, ?)
    """, (
        user_id,
        credentials.token,
        refresh_token,
        credentials.token_uri,
        expiry_str,
        json.dumps(credentials.scopes) if credentials.scopes else '[]'
    ))
    conn.commit()
    conn.close()

def get_calendar_service_for_user(user_id):
    """Get authenticated Calendar service for a specific user"""
    creds = get_user_calendar_credentials(user_id)
    if not creds:
        return None
    
    try:
        service = build('calendar', 'v3', credentials=creds)
        return service
    except Exception as e:
        print(f"Error building Calendar service: {e}")
        return None

def get_calendar_access_token():
    """Get Google Calendar access token from Replit connector"""
    hostname = os.environ.get('REPLIT_CONNECTORS_HOSTNAME')
    repl_identity = os.environ.get('REPL_IDENTITY')
    web_repl_renewal = os.environ.get('WEB_REPL_RENEWAL')
    
    x_replit_token = None
    if repl_identity:
        x_replit_token = 'repl ' + repl_identity
    elif web_repl_renewal:
        x_replit_token = 'depl ' + web_repl_renewal
    
    if not hostname or not x_replit_token:
        return None
    
    try:
        url = f'https://{hostname}/api/v2/connection?include_secrets=true&connector_names=google-calendar'
        headers = {
            'Accept': 'application/json',
            'X_REPLIT_TOKEN': x_replit_token
        }
        
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        
        data = response.json()
        items = data.get('items', [])
        
        if not items:
            return None
        
        connection_settings = items[0]
        access_token = connection_settings.get('settings', {}).get('access_token')
        
        if not access_token:
            access_token = connection_settings.get('settings', {}).get('oauth', {}).get('credentials', {}).get('access_token')
        
        return access_token
    except Exception as e:
        print(f"Error fetching calendar access token: {e}")
        return None

def get_calendar_service():
    """Get authenticated Google Calendar service using Replit connector"""
    access_token = get_calendar_access_token()
    
    if not access_token:
        return None
    
    try:
        from google.oauth2.credentials import Credentials as OAuth2Credentials
        creds = OAuth2Credentials(token=access_token)
        service = build('calendar', 'v3', credentials=creds)
        return service
    except Exception as e:
        print(f"Error building Calendar service: {e}")
        return None

def analyze_email_priority(subject, sender, snippet):
    """
    Use AI to analyze email priority
    Returns: 'urgent', 'important', or 'normal'
    
    Note: the newest OpenAI model is "gpt-5" which was released August 7, 2025.
    do not change this unless explicitly requested by the user
    """
    if not openai_client:
        return 'normal'
    
    try:
        prompt = f"""Analyze this email and categorize its priority level.

Subject: {subject}
From: {sender}
Preview: {snippet}

Categorize as one of:
- urgent: Requires immediate action (deadlines, urgent requests, time-sensitive)
- important: Significant but not time-critical (project updates, important decisions)
- normal: Regular communication (newsletters, routine updates, general info)

Respond with JSON in this exact format: {{"priority": "urgent|important|normal", "reason": "brief explanation"}}"""

        response = openai_client.chat.completions.create(
            model="gpt-5",
            messages=[
                {
                    "role": "system",
                    "content": "You are an email priority assistant. Analyze emails and categorize them accurately."
                },
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"},
            max_completion_tokens=200
        )
        
        result = json.loads(response.choices[0].message.content)
        priority = result.get('priority', 'normal').lower()
        
        if priority not in ['urgent', 'important', 'normal']:
            priority = 'normal'
        
        return priority
    except Exception as e:
        print(f"Error analyzing email priority: {e}")
        return 'normal'

def generate_email_replies(subject, sender, body):
    """
    Generate AI-powered email reply suggestions
    Returns: List of reply suggestions with different tones
    """
    if not openai_client:
        return []
    
    try:
        prompt = f"""Generate 3 professional email reply suggestions for this email with different tones:

Subject: {subject}
From: {sender}
Email Body:
{body[:1000]}

Generate 3 replies:
1. Professional - Formal and detailed
2. Friendly - Warm and conversational
3. Brief - Short and to the point

Each reply should:
- Address the email appropriately
- Be complete and ready to send
- Match the requested tone
- Be 2-4 sentences long

Respond with JSON in this exact format:
{{
  "replies": [
    {{"tone": "Professional", "text": "reply text here"}},
    {{"tone": "Friendly", "text": "reply text here"}},
    {{"tone": "Brief", "text": "reply text here"}}
  ]
}}"""

        response = openai_client.chat.completions.create(
            model="gpt-5",
            messages=[
                {
                    "role": "system",
                    "content": "You are a professional email assistant. Generate helpful, contextually appropriate email replies."
                },
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"},
            max_completion_tokens=500
        )
        
        result = json.loads(response.choices[0].message.content)
        return result.get('replies', [])
    except Exception as e:
        print(f"Error generating email replies: {e}")
        return []

def current_user():
    uid = session.get("user_id")
    if not uid:
        return None
    conn = get_db()
    row = conn.execute(
        "SELECT id, email, name, subscription_status, subscription_end_date, stripe_customer_id FROM users WHERE id = ?", 
        (uid,)
    ).fetchone()
    conn.close()
    # Convert sqlite3.Row to dict for easier access
    return dict(row) if row else None

def has_active_subscription(user):
    """Check if user has an active subscription"""
    if not user:
        return False
    
    status = user.get('subscription_status', 'free')
    if status == 'active':
        return True
    
    # Check if subscription end date is in the future
    end_date_str = user.get('subscription_end_date')
    if end_date_str:
        try:
            end_date = datetime.fromisoformat(end_date_str)
            return datetime.now() < end_date
        except:
            pass
    
    return False

def login_required(f):
    from functools import wraps
    @wraps(f)
    def wrapper(*args, **kwargs):
        if not current_user():
            flash("Please log in to continue.", "warning")
            return redirect(url_for("login", next=request.path))
        return f(*args, **kwargs)
    return wrapper

@app.route("/set-language/<lang>")
def set_language(lang):
    """Switch between English and Arabic"""
    if lang in ['en', 'ar']:
        session['language'] = lang
    return redirect(request.referrer or url_for('index'))

@app.route("/")
@login_required
def index():
    user = current_user()
    has_subscription = has_active_subscription(user)
    return render_template("index.html", user=user, has_subscription=has_subscription)

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")
        conn = get_db()
        row = conn.execute("SELECT id, password_hash, name FROM users WHERE email = ?", (email,)).fetchone()
        conn.close()
        if row and check_password_hash(row["password_hash"], password):
            session["user_id"] = row["id"]
            flash("Welcome back!", "success")
            nxt = request.args.get("next") or url_for("index")
            return redirect(nxt)
        flash("Invalid email or password.", "danger")
    return render_template("login.html")

@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")
        if not name or not email or not password:
            flash("All fields are required.", "warning")
            return render_template("register.html", name=name, email=email)
        pw_hash = generate_password_hash(password)
        conn = get_db()
        try:
            conn.execute("INSERT INTO users (email, name, password_hash) VALUES (?, ?, ?)", (email, name, pw_hash))
            conn.commit()
            flash("Account created. Please log in.", "success")
            return redirect(url_for("login"))
        except sqlite3.IntegrityError:
            flash("Email already registered.", "danger")
        finally:
            conn.close()
    return render_template("register.html")

@app.route("/logout")
def logout():
    session.clear()
    flash("Logged out.", "info")
    return redirect(url_for("login"))

@app.route("/pricing")
def pricing():
    """Show pricing plans"""
    return render_template("pricing.html")

@app.route("/create-checkout-session", methods=["POST"])
@login_required
def create_checkout_session():
    """Create a Stripe checkout session for subscription"""
    if not STRIPE_SECRET_KEY:
        flash("Payment system is not configured.", "danger")
        return redirect(url_for("pricing"))
    
    user = current_user()
    try:
        # Get domain for redirect URLs
        domain = os.environ.get('REPLIT_DEV_DOMAIN')
        if not domain:
            domains = os.environ.get('REPLIT_DOMAINS', '')
            if domains:
                domain = domains.split(',')[0]
        
        if not domain:
            domain = request.host
        
        # Create or retrieve Stripe customer
        customer_id = user.get('stripe_customer_id')
        if not customer_id:
            customer = stripe.Customer.create(
                email=user['email'],
                name=user['name'],
                metadata={'user_id': user['id']}
            )
            customer_id = customer.id
            
            # Save customer ID to database
            conn = get_db()
            conn.execute(
                "UPDATE users SET stripe_customer_id = ? WHERE id = ?",
                (customer_id, user['id'])
            )
            conn.commit()
            conn.close()
        
        # Create checkout session
        print(f"Creating Stripe checkout for user {user['id']} with domain {domain}")
        checkout_session = stripe.checkout.Session.create(
            customer=customer_id,
            line_items=[{
                'price_data': {
                    'currency': 'usd',
                    'unit_amount': 2900,  # $29.00 in cents
                    'recurring': {
                        'interval': 'month'
                    },
                    'product_data': {
                        'name': 'SmartOffice AI Premium',
                        'description': 'Full access to Email, Calendar, Reports & AI features',
                    },
                },
                'quantity': 1,
            }],
            mode='subscription',
            success_url=f'https://{domain}/success?session_id={{CHECKOUT_SESSION_ID}}',
            cancel_url=f'https://{domain}/cancel',
            metadata={'user_id': user['id']}
        )
        
        print(f"Stripe checkout session created: {checkout_session.id}")
        print(f"Redirect URL: {checkout_session.url}")
        
        # Store session URL in session for debugging
        session['stripe_checkout_url'] = checkout_session.url
        
        return redirect(checkout_session.url, code=303)
        
    except Exception as e:
        print(f"Error creating checkout session: {e}")
        flash(f"Payment error: {str(e)}", "danger")
        return redirect(url_for("pricing"))

@app.route("/success")
@login_required
def success():
    """Handle successful payment"""
    session_id = request.args.get('session_id')
    
    if session_id and STRIPE_SECRET_KEY:
        try:
            # Retrieve the session to confirm payment
            checkout_session = stripe.checkout.Session.retrieve(session_id)
            
            if checkout_session.payment_status == 'paid':
                user = current_user()
                # Calculate subscription end date (30 days from now)
                end_date = datetime.now() + timedelta(days=30)
                
                # Update user subscription status
                conn = get_db()
                conn.execute(
                    "UPDATE users SET subscription_status = ?, subscription_end_date = ? WHERE id = ?",
                    ('active', end_date.isoformat(), user['id'])
                )
                conn.commit()
                conn.close()
                
                flash("Subscription activated successfully!", "success")
        except Exception as e:
            print(f"Error processing successful payment: {e}")
    
    return render_template("success.html")

@app.route("/cancel")
def cancel():
    """Handle cancelled payment"""
    return render_template("cancel.html")

@app.route("/gmail-authorize")
@login_required
def gmail_authorize():
    """Start OAuth flow for Gmail"""
    user = current_user()
    
    if not GOOGLE_CLIENT_ID or not GOOGLE_CLIENT_SECRET:
        flash("Gmail integration is not configured. Please contact administrator.", "danger")
        return redirect(url_for('email'))
    
    redirect_uri = url_for('oauth2callback', _external=True, _scheme='https')
    
    try:
        flow = Flow.from_client_config(
            {
                "web": {
                    "client_id": GOOGLE_CLIENT_ID,
                    "client_secret": GOOGLE_CLIENT_SECRET,
                    "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                    "token_uri": "https://oauth2.googleapis.com/token",
                }
            },
            scopes=SCOPES,
            redirect_uri=redirect_uri
        )
        
        authorization_url, state = flow.authorization_url(
            access_type='offline',
            prompt='consent'
        )
        
        session['oauth_state'] = state
        return redirect(authorization_url)
    except Exception as e:
        print(f"[OAuth ERROR] Failed to start authorization: {e}")
        flash(f"OAuth configuration error: {str(e)}", "danger")
        return redirect(url_for('email'))

@app.route("/oauth2callback")
def oauth2callback():
    """Handle OAuth callback from Google"""
    user = current_user()
    if not user:
        flash("Please log in first.", "warning")
        return redirect(url_for('login'))
    
    state = session.get('oauth_state')
    redirect_uri = url_for('oauth2callback', _external=True, _scheme='https')
    
    flow = Flow.from_client_config(
        {
            "web": {
                "client_id": GOOGLE_CLIENT_ID,
                "client_secret": GOOGLE_CLIENT_SECRET,
                "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                "token_uri": "https://oauth2.googleapis.com/token",
            }
        },
        scopes=SCOPES,
        state=state,
        redirect_uri=redirect_uri
    )
    
    authorization_response = request.url.replace('http://', 'https://')
    
    try:
        flow.fetch_token(authorization_response=authorization_response)
        
        credentials = flow.credentials
        save_user_gmail_credentials(user['id'], credentials)
        
        # Return auto-close page that refreshes parent window
        return '''
        <html>
        <head><title>Gmail Connected</title></head>
        <body>
            <h3>Gmail connected successfully! Closing...</h3>
            <script>
                if (window.opener) {
                    window.opener.location.reload();
                    window.close();
                } else {
                    window.location.href = '/email';
                }
            </script>
        </body>
        </html>
        '''
    except Exception as e:
        flash(f"Failed to connect Gmail: {str(e)}", "danger")
        return redirect(url_for('email'))

@app.route("/calendar-authorize")
@login_required
def calendar_authorize():
    """Start OAuth flow for Google Calendar"""
    user = current_user()
    
    if not GOOGLE_CLIENT_ID or not GOOGLE_CLIENT_SECRET:
        flash("Calendar integration is not configured. Please contact administrator.", "danger")
        return redirect(url_for('calendar'))
    
    redirect_uri = url_for('calendar_oauth2callback', _external=True, _scheme='https')
    
    calendar_scopes = ['https://www.googleapis.com/auth/calendar.readonly']
    
    try:
        flow = Flow.from_client_config(
            {
                "web": {
                    "client_id": GOOGLE_CLIENT_ID,
                    "client_secret": GOOGLE_CLIENT_SECRET,
                    "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                    "token_uri": "https://oauth2.googleapis.com/token",
                }
            },
            scopes=calendar_scopes,
            redirect_uri=redirect_uri
        )
        
        authorization_url, state = flow.authorization_url(
            access_type='offline',
            prompt='consent'
        )
        
        session['calendar_oauth_state'] = state
        return redirect(authorization_url)
    except Exception as e:
        print(f"[Calendar OAuth ERROR] Failed to start authorization: {e}")
        flash(f"Calendar OAuth configuration error: {str(e)}", "danger")
        return redirect(url_for('calendar'))

@app.route("/calendar-oauth2callback")
def calendar_oauth2callback():
    """Handle OAuth callback from Google for Calendar"""
    user = current_user()
    if not user:
        flash("Please log in first.", "warning")
        return redirect(url_for('login'))
    
    state = session.get('calendar_oauth_state')
    redirect_uri = url_for('calendar_oauth2callback', _external=True, _scheme='https')
    
    calendar_scopes = ['https://www.googleapis.com/auth/calendar.readonly']
    
    flow = Flow.from_client_config(
        {
            "web": {
                "client_id": GOOGLE_CLIENT_ID,
                "client_secret": GOOGLE_CLIENT_SECRET,
                "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                "token_uri": "https://oauth2.googleapis.com/token",
            }
        },
        scopes=calendar_scopes,
        state=state,
        redirect_uri=redirect_uri
    )
    
    authorization_response = request.url.replace('http://', 'https://')
    
    try:
        flow.fetch_token(authorization_response=authorization_response)
        
        credentials = flow.credentials
        save_user_calendar_credentials(user['id'], credentials)
        
        # Return auto-close page that refreshes parent window
        return '''
        <html>
        <head><title>Calendar Connected</title></head>
        <body>
            <h3>Calendar connected successfully! Closing...</h3>
            <script>
                if (window.opener) {
                    window.opener.location.reload();
                    window.close();
                } else {
                    window.location.href = '/calendar';
                }
            </script>
        </body>
        </html>
        '''
    except Exception as e:
        flash(f"Failed to connect Calendar: {str(e)}", "danger")
        return redirect(url_for('calendar'))

@app.route("/email")
@login_required
def email():
    user = current_user()
    
    # Check subscription
    if not has_active_subscription(user):
        flash("Email features require a Premium subscription.", "warning")
        return redirect(url_for('pricing'))
    
    gmail_service = get_gmail_service(user['id'])
    
    if not gmail_service:
        return render_template("email.html", user=user, emails=[], connected=False)
    
    try:
        results = gmail_service.users().messages().list(userId='me', maxResults=15).execute()
        messages = results.get('messages', [])
        
        emails = []
        for msg in messages:
            msg_data = gmail_service.users().messages().get(userId='me', id=msg['id'], format='metadata', metadataHeaders=['From', 'Subject', 'Date']).execute()
            
            email_info = {
                'id': msg['id'],
                'subject': 'No Subject',
                'sender': 'Unknown',
                'date': 'Unknown',
                'snippet': msg_data.get('snippet', '')
            }
            
            for header in msg_data.get('payload', {}).get('headers', []):
                if header['name'] == 'Subject':
                    email_info['subject'] = header['value']
                elif header['name'] == 'From':
                    email_info['sender'] = header['value']
                elif header['name'] == 'Date':
                    email_info['date'] = header['value']
            
            # AI prioritization disabled until OpenAI credits added (prevents slow page loads)
            # When ready: Remove the next line and uncomment the block below
            email_info['priority'] = 'normal'
            
            # UNCOMMENT WHEN OPENAI CREDITS ACTIVE:
            # if openai_client and OPENAI_API_KEY:
            #     email_info['priority'] = analyze_email_priority(
            #         email_info['subject'],
            #         email_info['sender'],
            #         email_info['snippet']
            #     )
            # else:
            #     email_info['priority'] = 'normal'
            
            emails.append(email_info)
        
        return render_template("email.html", user=user, emails=emails, connected=True)
    
    except Exception as e:
        print(f"Error fetching emails: {str(e)}")
        flash(f"Error fetching emails: {str(e)}", "danger")
        return render_template("email.html", user=user, emails=[], connected=False)

@app.route("/email/<email_id>")
@login_required
def email_detail(email_id):
    """View full email with AI reply suggestions"""
    user = current_user()
    
    # Check subscription
    if not has_active_subscription(user):
        flash("Email features require a Premium subscription.", "warning")
        return redirect(url_for('pricing'))
    
    gmail_service = get_gmail_service(user['id'])
    
    if not gmail_service:
        flash("Gmail not connected.", "warning")
        return redirect(url_for('email'))
    
    try:
        msg_data = gmail_service.users().messages().get(
            userId='me', 
            id=email_id, 
            format='full'
        ).execute()
        
        email_info = {
            'id': email_id,
            'subject': 'No Subject',
            'sender': 'Unknown',
            'date': 'Unknown',
            'body': '',
            'snippet': msg_data.get('snippet', ''),
            'priority': 'normal'
        }
        
        for header in msg_data.get('payload', {}).get('headers', []):
            if header['name'] == 'Subject':
                email_info['subject'] = header['value']
            elif header['name'] == 'From':
                email_info['sender'] = header['value']
            elif header['name'] == 'Date':
                email_info['date'] = header['value']
        
        import base64
        import re
        from html import unescape
        
        def extract_email_body(payload):
            """Recursively extract email body, handling nested multipart structures"""
            body_text = ''
            body_html = ''
            
            def decode_part(data_str):
                if not data_str:
                    return ''
                try:
                    return base64.urlsafe_b64decode(data_str).decode('utf-8', errors='ignore')
                except:
                    return ''
            
            def strip_html(html_text):
                """Convert HTML to plain text"""
                text = re.sub(r'<style[^>]*>.*?</style>', '', html_text, flags=re.DOTALL | re.IGNORECASE)
                text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL | re.IGNORECASE)
                text = re.sub(r'<[^>]+>', ' ', text)
                text = unescape(text)
                text = re.sub(r'\s+', ' ', text)
                return text.strip()
            
            def parse_parts(parts):
                nonlocal body_text, body_html
                for part in parts:
                    mime_type = part.get('mimeType', '')
                    
                    if 'parts' in part:
                        parse_parts(part['parts'])
                    elif mime_type == 'text/plain':
                        data = part.get('body', {}).get('data', '')
                        decoded = decode_part(data)
                        if decoded and not body_text:
                            body_text = decoded
                    elif mime_type == 'text/html':
                        data = part.get('body', {}).get('data', '')
                        decoded = decode_part(data)
                        if decoded and not body_html:
                            body_html = decoded
            
            if 'parts' in payload:
                parse_parts(payload['parts'])
            else:
                mime_type = payload.get('mimeType', '')
                data = payload.get('body', {}).get('data', '')
                decoded = decode_part(data)
                
                if mime_type == 'text/plain':
                    body_text = decoded
                elif mime_type == 'text/html':
                    body_html = decoded
            
            if body_text:
                return body_text
            elif body_html:
                return strip_html(body_html)
            else:
                return ''
        
        payload = msg_data.get('payload', {})
        body = extract_email_body(payload)
        
        body_extracted_successfully = bool(body and body.strip())
        
        if not body_extracted_successfully:
            body = email_info['snippet'] + "\n\n[Note: Full email content could not be extracted. Showing preview only.]"
        
        email_info['body'] = body
        email_info['body_extracted'] = body_extracted_successfully
        
        reply_suggestions = []
        ai_enabled = bool(openai_client and OPENAI_API_KEY)
        
        if ai_enabled and body_extracted_successfully:
            reply_suggestions = generate_email_replies(
                email_info['subject'],
                email_info['sender'],
                body
            )
        
        return render_template(
            "email_detail.html",
            user=user,
            email=email_info,
            reply_suggestions=reply_suggestions,
            ai_enabled=ai_enabled
        )
    
    except Exception as e:
        print(f"Error fetching email detail: {str(e)}")
        flash(f"Error loading email: {str(e)}", "danger")
        return redirect(url_for('email'))

@app.route("/calendar")
@login_required
def calendar():
    user = current_user()
    
    # Check subscription
    if not has_active_subscription(user):
        flash("Calendar features require a Premium subscription.", "warning")
        return redirect(url_for('pricing'))
    
    # Use per-user calendar service instead of project-level connector
    calendar_service = get_calendar_service_for_user(user['id'])
    
    if not calendar_service:
        return render_template("calendar.html", user=user, events=[], connected=False)
    
    try:
        now = datetime.now(timezone.utc).isoformat()
        events_result = calendar_service.events().list(
            calendarId='primary',
            timeMin=now,
            maxResults=20,
            singleEvents=True,
            orderBy='startTime'
        ).execute()
        
        events = events_result.get('items', [])
        
        formatted_events = []
        for event in events:
            start = event['start'].get('dateTime', event['start'].get('date'))
            end = event['end'].get('dateTime', event['end'].get('date'))
            
            try:
                if 'T' in start:
                    start_dt = datetime.fromisoformat(start.replace('Z', '+00:00'))
                    start_formatted = start_dt.strftime('%b %d, %Y at %I:%M %p')
                else:
                    start_dt = datetime.fromisoformat(start)
                    start_formatted = start_dt.strftime('%b %d, %Y')
            except:
                start_formatted = start
            
            formatted_events.append({
                'id': event['id'],
                'summary': event.get('summary', 'No Title'),
                'start': start_formatted,
                'description': event.get('description', ''),
                'location': event.get('location', ''),
                'htmlLink': event.get('htmlLink', '')
            })
        
        return render_template("calendar.html", user=user, events=formatted_events, connected=True)
    
    except Exception as e:
        print(f"Error fetching calendar events: {str(e)}")
        flash(f"Error fetching calendar events: {str(e)}", "danger")
        return render_template("calendar.html", user=user, events=[], connected=False)

@app.route("/reports")
@login_required
def reports():
    user = current_user()
    
    # Check subscription
    if not has_active_subscription(user):
        flash("Report features require a Premium subscription.", "warning")
        return redirect(url_for('pricing'))
    
    gmail_service = get_gmail_service(user['id'])
    calendar_service = get_calendar_service_for_user(user['id'])
    
    gmail_connected = bool(gmail_service)
    calendar_connected = bool(calendar_service)
    
    return render_template("reports.html", user=user, 
                         gmail_connected=gmail_connected,
                         calendar_connected=calendar_connected)

@app.route("/generate-report", methods=["POST"])
@login_required
def generate_report():
    user = current_user()
    
    report_period = request.form.get('report_period')
    report_format = request.form.get('report_format', 'pdf')
    include_emails = 'include_emails' in request.form
    include_calendar = 'include_calendar' in request.form
    
    if report_period == 'week':
        start_date = datetime.now() - timedelta(days=7)
        end_date = datetime.now()
        period_name = "Past Week"
    elif report_period == 'month':
        start_date = datetime.now() - timedelta(days=30)
        end_date = datetime.now()
        period_name = "Past Month"
    else:
        start_date_str = request.form.get('start_date')
        end_date_str = request.form.get('end_date')
        if not start_date_str or not end_date_str:
            flash("Please provide start and end dates for custom period.", "danger")
            return redirect(url_for('reports'))
        start_date = datetime.strptime(start_date_str, '%Y-%m-%d')
        end_date = datetime.strptime(end_date_str, '%Y-%m-%d')
        period_name = f"{start_date.strftime('%b %d')} - {end_date.strftime('%b %d, %Y')}"
    
    report_data = {
        'user_name': user['name'],
        'period': period_name,
        'start_date': start_date,
        'end_date': end_date,
        'emails': [],
        'events': []
    }
    
    if include_emails:
        gmail_service = get_gmail_service(user['id'])
        if gmail_service:
            try:
                after_timestamp = int(start_date.timestamp())
                results = gmail_service.users().messages().list(
                    userId='me',
                    maxResults=50,
                    q=f'after:{after_timestamp}'
                ).execute()
                messages = results.get('messages', [])
                
                for msg in messages[:20]:
                    msg_data = gmail_service.users().messages().get(
                        userId='me',
                        id=msg['id'],
                        format='metadata',
                        metadataHeaders=['From', 'Subject', 'Date']
                    ).execute()
                    
                    email_info = {'subject': 'No Subject', 'sender': 'Unknown', 'date': 'Unknown'}
                    for header in msg_data.get('payload', {}).get('headers', []):
                        if header['name'] == 'Subject':
                            email_info['subject'] = header['value']
                        elif header['name'] == 'From':
                            email_info['sender'] = header['value']
                        elif header['name'] == 'Date':
                            email_info['date'] = header['value']
                    
                    report_data['emails'].append(email_info)
            except Exception as e:
                print(f"Error fetching emails for report: {e}")
    
    if include_calendar:
        calendar_service = get_calendar_service_for_user(user['id'])
        if calendar_service:
            try:
                time_min = start_date.isoformat() + 'Z'
                time_max = end_date.isoformat() + 'Z'
                
                events_result = calendar_service.events().list(
                    calendarId='primary',
                    timeMin=time_min,
                    timeMax=time_max,
                    maxResults=50,
                    singleEvents=True,
                    orderBy='startTime'
                ).execute()
                
                events = events_result.get('items', [])
                for event in events:
                    start = event['start'].get('dateTime', event['start'].get('date'))
                    report_data['events'].append({
                        'summary': event.get('summary', 'No Title'),
                        'start': start,
                        'location': event.get('location', '')
                    })
            except Exception as e:
                print(f"Error fetching calendar events for report: {e}")
    
    if report_format == 'pdf':
        return generate_pdf_report(report_data)
    else:
        return generate_docx_report(report_data)

def generate_pdf_report(data):
    """Generate PDF report using ReportLab"""
    from xml.sax.saxutils import escape
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75*inch, bottomMargin=0.75*inch)
    story = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#0d6efd'),
        spaceAfter=30,
        alignment=1
    )
    
    story.append(Paragraph("SmartOffice AI Report", title_style))
    story.append(Paragraph(f"<b>For:</b> {escape(data['user_name'])}", styles['Normal']))
    story.append(Paragraph(f"<b>Period:</b> {escape(data['period'])}", styles['Normal']))
    story.append(Spacer(1, 0.3*inch))
    
    if data['emails']:
        story.append(Paragraph(f"<b>Email Summary ({len(data['emails'])} emails)</b>", styles['Heading2']))
        story.append(Spacer(1, 0.2*inch))
        
        email_data = [['Date', 'From', 'Subject']]
        for email in data['emails'][:15]:
            email_data.append([
                Paragraph(escape(email['date'][:20]), styles['Normal']),
                Paragraph(escape(email['sender'][:30]), styles['Normal']),
                Paragraph(escape(email['subject'][:50]), styles['Normal'])
            ])
        
        table = Table(email_data, colWidths=[1.5*inch, 2*inch, 3*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0d6efd')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ]))
        story.append(table)
        story.append(Spacer(1, 0.4*inch))
    
    if data['events']:
        story.append(Paragraph(f"<b>Calendar Events ({len(data['events'])} events)</b>", styles['Heading2']))
        story.append(Spacer(1, 0.2*inch))
        
        from xml.sax.saxutils import escape
        for event in data['events'][:15]:
            story.append(Paragraph(f"• <b>{escape(event['summary'])}</b>", styles['Normal']))
            story.append(Paragraph(f"  {escape(event['start'])} {escape(event['location'])}", styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
    
    doc.build(story)
    buffer.seek(0)
    
    filename = f"SmartOffice_Report_{datetime.now().strftime('%Y%m%d')}.pdf"
    return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/pdf')

def generate_docx_report(data):
    """Generate Word document report"""
    doc = Document()
    
    title = doc.add_heading('SmartOffice AI Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"For: {data['user_name']}")
    doc.add_paragraph(f"Period: {data['period']}")
    doc.add_paragraph()
    
    if data['emails']:
        doc.add_heading(f'Email Summary ({len(data["emails"])} emails)', level=1)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Date'
        hdr_cells[1].text = 'From'
        hdr_cells[2].text = 'Subject'
        
        for email in data['emails'][:15]:
            row_cells = table.add_row().cells
            row_cells[0].text = email['date'][:20]
            row_cells[1].text = email['sender'][:30]
            row_cells[2].text = email['subject'][:50]
        
        doc.add_paragraph()
    
    if data['events']:
        doc.add_heading(f'Calendar Events ({len(data["events"])} events)', level=1)
        
        for event in data['events'][:15]:
            doc.add_paragraph(f"{event['summary']}", style='List Bullet')
            doc.add_paragraph(f"  {event['start']} {event['location']}", style='Body Text')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"SmartOffice_Report_{datetime.now().strftime('%Y%m%d')}.docx"
    return send_file(buffer, as_attachment=True, download_name=filename, 
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route("/privacy")
def privacy():
    """Privacy Policy page"""
    return render_template('privacy.html')

@app.route("/terms")
def terms():
    """Terms of Service page"""
    return render_template('terms.html')

@app.route("/download-cost-breakdown")
def download_cost_breakdown():
    """Generate Word document with cost breakdown"""
    doc = Document()
    
    # Title
    title = doc.add_heading('SmartOffice AI - Cost Breakdown & Requirements', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Last Updated: October 8, 2025')
    doc.add_paragraph()
    
    # Section 1: Actual Costs
    doc.add_heading('1. Actual Monthly Costs Breakdown', level=1)
    
    doc.add_heading('1.1 Replit Deployment (REQUIRED)', level=2)
    doc.add_paragraph('Cost: $3-5/month for small apps')
    p = doc.add_paragraph()
    p.add_run('  • $1.00 base fee/month\n')
    p.add_run('  • Plus usage (scales with traffic)\n')
    p.add_run('  • Example: 500 visitors/day = ~$3/month total')
    doc.add_paragraph('What you get: Hosting, SSL, automatic scaling')
    doc.add_paragraph('Status: ✅ Already configured')
    doc.add_paragraph()
    
    doc.add_heading('1.2 Google Cloud (FREE - NO COST!)', level=2)
    doc.add_paragraph('Cost: $0 (FREE for OAuth and APIs)')
    doc.add_paragraph('Free tier includes:')
    p = doc.add_paragraph()
    p.add_run('  • Unlimited OAuth requests\n')
    p.add_run('  • Gmail API: 1 billion quota units/day\n')
    p.add_run('  • Calendar API: 1 million requests/day')
    doc.add_paragraph('You only pay if you exceed limits (unlikely for SME use)')
    doc.add_paragraph('Status: ✅ Already set up')
    doc.add_paragraph()
    
    doc.add_heading('1.3 OpenAI API (PAY-AS-YOU-GO)', level=2)
    doc.add_paragraph('Cost: $20-50/month depending on usage')
    p = doc.add_paragraph()
    p.add_run('  • GPT-4o: $2.50 per 1M input tokens\n')
    p.add_run('  • NOT a subscription - pay only for what you use\n')
    p.add_run('  • NOT ChatGPT Plus ($20/month) - Different service!')
    doc.add_paragraph('Usage estimate: 100 emails analyzed/day ≈ $15-30/month')
    doc.add_paragraph('Status: ✅ Already integrated, needs API credits')
    doc.add_paragraph()
    
    doc.add_heading('1.4 Stripe Payments (TRANSACTION FEE ONLY)', level=2)
    doc.add_paragraph('Monthly cost: $0')
    doc.add_paragraph('Per transaction: 2.9% + $0.30')
    p = doc.add_paragraph()
    p.add_run('  • Example: $29 subscription = $1.14 fee to Stripe\n')
    p.add_run('  • You keep: $27.86')
    doc.add_paragraph('No monthly subscription needed')
    doc.add_paragraph('Status: ✅ Already integrated')
    doc.add_paragraph()
    
    # Section 2: Cost Summary
    doc.add_heading('2. Total Monthly Costs Summary', level=1)
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Service'
    hdr_cells[1].text = 'Cost'
    hdr_cells[2].text = 'Notes'
    
    table.rows[1].cells[0].text = 'Replit Hosting'
    table.rows[1].cells[1].text = '$3-5'
    table.rows[1].cells[2].text = 'Required'
    
    table.rows[2].cells[0].text = 'Google Cloud'
    table.rows[2].cells[1].text = '$0'
    table.rows[2].cells[2].text = 'FREE tier sufficient'
    
    table.rows[3].cells[0].text = 'OpenAI API'
    table.rows[3].cells[1].text = '$20-50'
    table.rows[3].cells[2].text = 'Pay-as-you-go'
    
    table.rows[4].cells[0].text = 'Stripe'
    table.rows[4].cells[1].text = '$0 + 2.9% fees'
    table.rows[4].cells[2].text = 'Only on transactions'
    
    table.rows[5].cells[0].text = 'TOTAL'
    table.rows[5].cells[1].text = '$23-55/month'
    table.rows[5].cells[2].text = 'Covers everything'
    
    doc.add_paragraph()
    
    # Section 3: What You DON'T Need
    doc.add_heading('3. What You DON\'T Need', level=1)
    p = doc.add_paragraph()
    p.add_run('❌ Google Cloud paid subscription\n')
    p.add_run('❌ ChatGPT Plus ($20/month) - This is different from API\n')
    p.add_run('❌ Stripe monthly subscription\n')
    p.add_run('❌ "GPT-5" - You\'re using GPT-4o API')
    doc.add_paragraph()
    
    # Section 4: Revenue Math
    doc.add_heading('4. Revenue Math Example', level=1)
    doc.add_paragraph('If you get 10 customers at $29/month:')
    p = doc.add_paragraph()
    p.add_run('  • Revenue: $290/month\n')
    p.add_run('  • Costs: ~$50/month (Replit + OpenAI + Stripe fees)\n')
    p.add_run('  • Profit: ~$240/month 💰')
    doc.add_paragraph()
    doc.add_paragraph('Break-even: Just 2 customers!')
    doc.add_paragraph()
    
    # Section 5: To Start
    doc.add_heading('5. Steps to Get Started', level=1)
    p = doc.add_paragraph()
    p.add_run('1. Add $20 to OpenAI API (platform.openai.com/account/billing)\n')
    p.add_run('2. Deploy on Replit ($3-5/month, auto-charged)\n')
    p.add_run('3. That\'s it! Google and Stripe are free to start')
    doc.add_paragraph()
    
    # Section 6: Google Verification
    doc.add_heading('6. Google App Verification Process', level=1)
    
    doc.add_heading('Step 1: Access Google Cloud Console', level=2)
    p = doc.add_paragraph()
    p.add_run('1. Go to https://console.cloud.google.com\n')
    p.add_run('2. Select your SmartOffice AI project')
    doc.add_paragraph()
    
    doc.add_heading('Step 2: Navigate to OAuth Consent Screen', level=2)
    p = doc.add_paragraph()
    p.add_run('1. Click "APIs & Services" in left sidebar\n')
    p.add_run('2. Click "OAuth consent screen"')
    doc.add_paragraph()
    
    doc.add_heading('Step 3: Submit for Verification', level=2)
    doc.add_paragraph('Click the "Submit for Verification" button')
    doc.add_paragraph()
    
    doc.add_heading('Step 4: Required Information', level=2)
    doc.add_paragraph('App Information:')
    p = doc.add_paragraph()
    p.add_run('  • App Name: SmartOffice AI\n')
    p.add_run('  • App Description: "SmartOffice AI is an intelligent administrative assistant that helps SMEs organize emails, manage calendars, and generate automated reports using AI technology."')
    doc.add_paragraph()
    
    doc.add_paragraph('Legal Documents (use your Replit deployment URL):')
    p = doc.add_paragraph()
    p.add_run('  • Privacy Policy URL: https://your-replit-url.repl.co/privacy\n')
    p.add_run('  • Terms of Service URL: https://your-replit-url.repl.co/terms\n')
    p.add_run('  • Homepage URL: https://your-replit-url.repl.co')
    doc.add_paragraph()
    
    doc.add_paragraph('Scope Justification:')
    doc.add_paragraph('Gmail Scope (gmail.readonly): "We need read-only access to user emails to provide AI-powered priority categorization, smart reply suggestions, and email summary reports. No emails are stored permanently; they are only processed in real-time for analysis."')
    doc.add_paragraph()
    doc.add_paragraph('Calendar Scope (calendar.readonly): "We need read-only access to view calendar events to display upcoming meetings and include them in automated weekly reports. No calendar data is modified or stored."')
    doc.add_paragraph()
    
    doc.add_heading('Step 5: Review Timeline', level=2)
    p = doc.add_paragraph()
    p.add_run('  • Submit all information\n')
    p.add_run('  • Google reviews in 3-7 business days\n')
    p.add_run('  • Receive email updates at your Google account\n')
    p.add_run('  • After approval: No "unverified app" warning!')
    
    # Generate file
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = "SmartOffice_Cost_Breakdown.docx"
    return send_file(buffer, as_attachment=True, download_name=filename, 
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)), debug=True)
